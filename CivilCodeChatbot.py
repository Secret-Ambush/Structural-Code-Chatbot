import os
import base64
from dotenv import load_dotenv
import streamlit as st
from streamlit_chat import message
import fitz
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.agents import initialize_agent, Tool
from langchain.agents.agent_types import AgentType
from langchain.tools import Tool as LangchainTool
from langchain.utilities import SerpAPIWrapper
import tempfile
import re

load_dotenv()

st.set_page_config(page_title="Structural Code Chatbot", layout="wide")
page_bg_img = '''
<style>
[data-testid="stAppViewContainer"] {
    background-image: url("https://a.l3n.co/i/ovFKWD.png");
    background-size: cover;
    background-position: center;
    background-repeat: no-repeat;
    background-attachment: fixed;
}
</style>
'''
st.markdown(page_bg_img, unsafe_allow_html=True)
st.title("📘 Structural Code Chatbot for Civil Engineers")
st.markdown("This assistant helps you navigate structural engineering design codes like ACI 318, Eurocode, IS 456 and more. You can ask clause-related questions, run design checks, or calculate structural capacity using uploaded code documents.")

if st.button("🧹 Clear Chat"):
    st.session_state.messages = []
    st.rerun()

@st.cache_data(show_spinner=False)
def parse_pdf(file_path):
    doc = fitz.open(file_path)
    return "\n".join([page.get_text() for page in doc])

def parse_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def parse_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def load_local_documents(data_folder="data"):
    docs = []
    for fname in os.listdir(data_folder):
        file_path = os.path.join(data_folder, fname)
        suffix = fname.split(".")[-1].lower()
        if suffix not in ["pdf", "docx", "txt"]:
            continue
        try:
            if suffix == "pdf":
                text = parse_pdf(file_path)
            elif suffix == "docx":
                text = parse_docx(file_path)
            elif suffix == "txt":
                text = parse_txt(file_path)
            else:
                continue
            docs.append(Document(page_content=text, metadata={"source": fname}))
        except Exception as e:
            st.warning(f"Failed to load {fname}: {str(e)}")
    return docs

st.sidebar.markdown("### 📎 Upload Building Codes")
uploaded_files = st.sidebar.file_uploader("Upload PDFs, DOCX or TXT files", type=["pdf", "docx", "txt"], accept_multiple_files=True)

if "cached_documents" not in st.session_state:
    st.session_state.cached_documents = []

if uploaded_files:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    for uploaded_file in uploaded_files:
        suffix = uploaded_file.name.split(".")[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
            tmp.write(uploaded_file.getbuffer())
            tmp_path = tmp.name
        if suffix == "pdf":
            doc = fitz.open(tmp_path)
            for i, page in enumerate(doc):
                chunk = page.get_text()
                st.session_state.cached_documents.append(Document(page_content=chunk, metadata={"source": f"{uploaded_file.name} - Page {i+1}"}))
        elif suffix == "docx":
            text = parse_docx(tmp_path)
            for chunk in splitter.split_text(text):
                st.session_state.cached_documents.append(Document(page_content=chunk, metadata={"source": uploaded_file.name}))
        elif suffix == "txt":
            text = parse_txt(tmp_path)
            for chunk in splitter.split_text(text):
                st.session_state.cached_documents.append(Document(page_content=chunk, metadata={"source": uploaded_file.name}))

if st.session_state.cached_documents:
    st.sidebar.markdown("**Uploaded Files:**")
    uploaded_names = list({doc.metadata['source'] for doc in st.session_state.cached_documents})
    selected_file = st.sidebar.selectbox("Select a file to explore", uploaded_names)
    preview_text = next((doc.page_content[:500] for doc in st.session_state.cached_documents if doc.metadata['source'] == selected_file), "No preview available.")

    if selected_file.endswith(".pdf"):
        show_preview = st.sidebar.checkbox("Show Preview (PDF)")
        if show_preview:
            st.sidebar.markdown(f"Previewing: **{selected_file}**")
            for uploaded_file in uploaded_files:
                if uploaded_file.name == selected_file:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(uploaded_file.getbuffer())
                        pdf_preview_path = tmp.name
                    if os.path.exists(pdf_preview_path):
                        with open(pdf_preview_path, "rb") as f:
                            base64_pdf = base64.b64encode(f.read()).decode("utf-8")
                        pdf_display = f"""
                        <iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="400" type="application/pdf"></iframe>
                        """
                        st.sidebar.markdown(pdf_display, unsafe_allow_html=True)
    elif selected_file.endswith(".docx"):
        show_docx = st.sidebar.checkbox("Show DOCX Preview")
        if show_docx:
            st.sidebar.markdown(f"Previewing: **{selected_file}**")
            st.sidebar.text_area("Preview", preview_text, height=200)
    else:
        st.sidebar.markdown(f"Previewing: **{selected_file}**")
        st.sidebar.text_area("Preview", preview_text, height=200)

if "messages" not in st.session_state:
    st.session_state.messages = []

for i, msg in enumerate(st.session_state.messages):
    message(msg["content"], is_user=(msg["role"] == "user"), key=str(i))

user_query = st.chat_input("Ask about structural design codes...")

if user_query:
    st.session_state.messages.append({"role": "user", "content": user_query})
    placeholder = st.empty()
    placeholder.chat_message("assistant").write("...")

    local_docs = load_local_documents()
    if "retriever" not in st.session_state:
        documents = st.session_state.cached_documents + local_docs

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        if local_docs:
            chunked_local = splitter.split_documents(local_docs)
            documents.extend(chunked_local)

        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        vectorstore = FAISS.from_documents(documents, embeddings)
        st.session_state.retriever = vectorstore.as_retriever()

    retriever = st.session_state.retriever

    def calc_moment_capacity(query: str) -> str:
        numbers = list(map(float, re.findall(r"\d+\.?\d*", query)))
        if "shear" in query.lower() and len(numbers) >= 3:
            b, d, fc = numbers[:3]
            phi = 0.75
            Vc = 0.17 * (fc ** 0.5) * b * d
            phi_Vc = phi * Vc / 1000
            return f"🧮 ϕVc = {phi_Vc:.2f} kN (Nominal shear strength = {Vc/1000:.2f} kN)\nFormula: ϕVc = 0.75 × 0.17 × √fc × b × d\nClause: ACI 318-19 §22.5.5.1"
        elif "min reinforcement" in query.lower() and len(numbers) >= 2:
            b, d = numbers[:2]
            As_min = 0.0018 * b * d
            return f"📏 Minimum As = {As_min:.2f} mm²\nFormula: As,min = 0.0018 × b × d\nClause: ACI 318-19 §9.6.1.2"
        elif "load capacity" in query.lower() and len(numbers) >= 3:
            b, d, fc = numbers[:3]
            Pu = 0.4 * fc * b * d
            return f"⚖️ Factored axial load capacity ≈ {Pu/1000:.2f} kN\nFormula: Pu = 0.4 × fc × b × d\nClause: Approximate"
        elif "deflection" in query.lower() and len(numbers) >= 1:
            span = numbers[0]
            limit = span / 20
            return f"📐 Max allowable deflection ≈ L/20 = {limit:.1f} mm\nFormula: Δ_max = span / 20\nClause: ACI 318-19 Table 24.2.2"
        elif len(numbers) >= 4:
            b, d, fc, As = numbers[:4]
            phi = 0.9
            a = As * 420 / (0.85 * fc * b)
            Mn = As * 420 * (d - a / 2) / 1e6
            return f"🧮 ϕMn = {phi * Mn:.2f} kNm (Nominal moment = {Mn:.2f} kNm)\nFormula: ϕMn = 0.9 × As × fy × (d - a/2)\nClause: ACI 318-19 §22.3"
        return "❌ Could not determine the type of calculation or extract enough parameters. Try including keywords like 'moment', 'shear', 'min reinforcement', or 'deflection'."

    search = SerpAPIWrapper()

    tools = [
        Tool(name="Code QA", func=lambda q: RetrievalQA.from_chain_type(llm=ChatOpenAI(temperature=0), retriever=retriever).run(q), description="Search structural code documents."),
        Tool(name="Moment Capacity Calculator", func=calc_moment_capacity, description="Calculate structural values such as ϕMn, ϕVc, As,min, Pu, deflection."),
        LangchainTool(name="Web Search", func=search.run, description="Search the web for structural engineering or design-related information.")
    ]

    agent = initialize_agent(tools, ChatOpenAI(temperature=0), agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION, verbose=False)

    with st.spinner("Thinking..."):
        casual_phrases = ["hello", "hi", "hey", "what can you do", "who are you"]
        if any(phrase in user_query.lower() for phrase in casual_phrases):
            response = "👋 Hello! I'm your assistant for civil and structural engineering codes. You can ask me about ACI clauses, beam design, or code compliance checks."
        else:
            llm_check = ChatOpenAI(temperature=0)
            check_prompt = f"Is the following query related to civil or structural engineering design codes? Answer 'Yes' or 'No'.\nQuery: {user_query}"
            is_structural = llm_check.predict(check_prompt).strip().lower()
            if "yes" in is_structural:
                response = agent.run(user_query)
            else:
                response = "I'm designed to answer structural engineering and design code queries. Please ask something related to that. 🏗️"

    placeholder.empty()
    st.session_state.messages.append({"role": "assistant", "content": response})
else:
    st.warning("No documents found. Please upload files or check the data folder.")
